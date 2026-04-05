"""
Resume Parser
Takes raw file bytes → returns a filled ResumeProfile
"""
import io
import re
import uuid
import logging
from pathlib import Path
from app.services.skill_extractor import extract_skills
import pdfplumber                    # reads PDF files
from docx import Document as DocxDocument  # reads DOCX files

from app.models.schemas import ResumeProfile, ExperienceEntry, EducationEntry

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".jpg", ".jpeg", ".png"}

# Keywords that signal the start of each resume section
SECTION_KEYWORDS = {
    "experience": [
        "work experience", "professional experience", "employment",
        "career history", "experience", "work history"
    ],
    "education": [
        "education", "academic background", "qualifications", "degrees"
    ],
    "skills": [
        "skills", "technical skills", "core competencies",
        "expertise", "technologies", "tools"
    ],
    "projects": [
        "projects", "personal projects", "academic projects", "portfolio"
    ],
    "certifications": [
        "certifications", "certificates", "licenses"
    ],
    "summary": [
        "summary", "profile", "objective", "about me", "professional summary"
    ],
}

# Patterns to find years of experience mentioned in text
# e.g. "5 years of experience", "3+ years", "over 2 years"
EXPERIENCE_PATTERNS = [
    r"(\d+)\+?\s*years?\s*of\s*experience",
    r"(\d+)\+?\s*yrs?\s*(?:of\s*)?experience",
    r"over\s*(\d+)\s*years?",
    r"more\s*than\s*(\d+)\s*years?",
]


# ─────────────────────────────────────────────────────────────────────────────
# STEP 1: Extract raw text from file
# ─────────────────────────────────────────────────────────────────────────────

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Use pdfplumber to pull all text out of a PDF."""
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            # x_tolerance/y_tolerance help handle two-column layouts
            text = page.extract_text(x_tolerance=3, y_tolerance=3)
            if text:
                text_parts.append(text)
    return "\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from every paragraph and table in a DOCX."""
    doc = DocxDocument(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also grab text from tables (some resumes use table layouts)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def extract_raw_text(file_bytes: bytes, filename: str) -> str:
    """Route to the right extractor based on file extension."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_bytes)
    return ""   # images need OCR — we'll skip for now


# ─────────────────────────────────────────────────────────────────────────────
# STEP 2: Split text into sections
# ─────────────────────────────────────────────────────────────────────────────

def detect_sections(text: str) -> dict[str, str]:
    """
    Walk through each line. When a line matches a section heading keyword,
    start collecting text under that section label.
    Returns: { "experience": "...", "education": "...", ... }
    """
    lines = text.splitlines()
    sections: dict[str, str] = {"unknown": ""}
    current_section = "unknown"

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        lower = stripped.lower()
        matched = False

        # Check if this line looks like a section heading
        for section_name, keywords in SECTION_KEYWORDS.items():
            if any(kw in lower for kw in keywords) and len(stripped) < 60:
                current_section = section_name
                sections.setdefault(current_section, "")
                matched = True
                break

        if not matched:
            sections.setdefault(current_section, "")
            sections[current_section] += stripped + "\n"

    return {k: v.strip() for k, v in sections.items() if v.strip()}


# ─────────────────────────────────────────────────────────────────────────────
# STEP 3: Extract specific fields
# ─────────────────────────────────────────────────────────────────────────────

def extract_email(text: str) -> str:
    match = re.search(r"[\w.+-]+@[\w-]+\.[a-zA-Z]{2,}", text)
    return match.group(0) if match else ""


def extract_phone(text: str) -> str:
    match = re.search(r"[\+\(]?\d[\d\s\-\(\)]{7,}\d", text)
    return match.group(0).strip() if match else ""


def extract_name(text: str) -> str:
    """
    Heuristic: the candidate's name is usually one of the first lines.
    We look for a line with 2-4 words, all capitalised, no digits or symbols.
    """
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    for line in lines[:6]:   # only check first 6 lines
        words = line.split()
        if (2 <= len(words) <= 4
                and all(w[0].isupper() for w in words if w)
                and not any(c.isdigit() for c in line)
                and "@" not in line
                and len(line) < 50):
            return line
    return "Unknown"


def extract_total_experience(text: str) -> float:
    """Find the highest number of years mentioned in the text."""
    years = []
    for pattern in EXPERIENCE_PATTERNS:
        for match in re.finditer(pattern, text, re.IGNORECASE):
            try:
                years.append(float(match.group(1)))
            except (IndexError, ValueError):
                pass
    return max(years) if years else 0.0


def extract_education(section_text: str) -> list[EducationEntry]:
    """Look for degree keywords and extract education entries."""
    degree_keywords = [
        "bachelor", "master", "phd", "doctorate",
        "b.tech", "m.tech", "b.sc", "m.sc", "mba", "b.e", "m.e"
    ]
    entries = []
    blocks = re.split(r"\n{2,}", section_text)

    for block in blocks:
        if any(kw in block.lower() for kw in degree_keywords):
            lines = [l.strip() for l in block.splitlines() if l.strip()]
            degree = lines[0] if lines else ""
            institution = lines[1] if len(lines) > 1 else ""
            year_match = re.search(r"\b(19|20)\d{2}\b", block)
            year = int(year_match.group()) if year_match else None
            entries.append(EducationEntry(
                degree=degree,
                institution=institution,
                year=year,
            ))
    return entries


def extract_certifications(section_text: str) -> list[str]:
    lines = [l.strip() for l in section_text.splitlines() if l.strip()]
    return lines[:10]


# ─────────────────────────────────────────────────────────────────────────────
# STEP 4: Put it all together
# ─────────────────────────────────────────────────────────────────────────────

def parse_resume(
    file_bytes: bytes,
    filename: str,
    job_role: str = "",
    upload_batch: str = "",
) -> ResumeProfile:
    """
    Main function — call this with raw file bytes, get back a ResumeProfile.
    """
    ext = Path(filename).suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        return ResumeProfile(
            id=str(uuid.uuid4()),
            filename=filename,
            parse_status="unsupported",
            job_role=job_role,
            upload_batch=upload_batch,
        )

    try:
        # Step 1: get raw text
        raw_text = extract_raw_text(file_bytes, filename)

        if not raw_text.strip():
            return ResumeProfile(
                id=str(uuid.uuid4()),
                filename=filename,
                parse_status="empty — could not extract text",
                job_role=job_role,
                upload_batch=upload_batch,
            )

        # Step 2: split into sections
        sections = detect_sections(raw_text)
        all_skills = []
        for section_name, section_text in sections.items():
            all_skills.extend(extract_skills(section_text, section_name))

        # Deduplicate — same skill might appear in multiple sections
        seen_skills: set[str] = set()
        unique_skills = []
        for skill in all_skills:
            if skill.canonical not in seen_skills:
                seen_skills.add(skill.canonical)
                unique_skills.append(skill)

        # Step 3: extract fields
        name  = extract_name(raw_text)
        email = extract_email(raw_text)
        phone = extract_phone(raw_text)

        experience_text = sections.get("experience", "")
        total_exp = extract_total_experience(experience_text or raw_text)
        if total_exp == 0 and "projects" in sections:
            project_count = sections["projects"].count("•")
            total_exp = round(project_count * 0.3, 1)
        education = extract_education(sections.get("education", ""))
        certs     = extract_certifications(sections.get("certifications", ""))
        summary   = sections.get("summary", "")

        # Step 4: return filled profile
        return ResumeProfile(
            id=str(uuid.uuid4()),
            filename=filename,
            candidate_name=name,
            email=email,
            phone=phone,
            total_experience_years=round(total_exp, 1),
            skills=unique_skills,
            education_entries=education,
            certifications=certs,
            summary=summary,
            raw_sections=sections,
            parse_status="success",
            job_role=job_role,
            upload_batch=upload_batch,
        )

    except Exception as exc:
        logger.exception("Failed to parse %s", filename)
        return ResumeProfile(
            id=str(uuid.uuid4()),
            filename=filename,
            parse_status=f"error: {str(exc)[:100]}",
            job_role=job_role,
            upload_batch=upload_batch,
        )